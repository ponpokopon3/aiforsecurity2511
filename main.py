#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
main2.py

目的:
  セキュリティ規則一覧ルールをチェックする CUI ベースの AgentRAG チャットボット

要件（主なポイント）:
  - 起動時に `rule/` の json ファイルを再帰的に読み込む
  - 起動時に `specification/` の pdf/docx/md/txt を読み込み、ChromaDB に格納する
  - Agent A: ドキュメント要約エージェント
  - Agent B: ドキュメント確認エージェント
    - Agent C/D は不要になったため本実装では Agent A/B に集中します
  - LCEL（Runnable 等）を利用してチェーンを組み立てる
  - OpenAI の gpt-4o を ChatOpenAI で呼び出す（OPENAI_API_KEY を利用）

注記:
  - 可読性優先・日本語コメント多め
  - メモリ節約: ドキュメントはチャンクし、検索時は上位 k のみを使用
  - 依存パッケージは README にも書いていますが、`requirements.txt` を参照してください
"""

import os
import sys
import json
import glob
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import textwrap
from datetime import datetime

# --- 依存パッケージのインポート（利用環境でインストールされている前提） ---
try:
    # LangChain とモデルラッパー
    from langchain.chat_models import ChatOpenAI
    from langchain.schema import HumanMessage, SystemMessage
    from langchain.embeddings import HuggingFaceEmbeddings
    from langchain.vectorstores import Chroma
    from langchain.text_splitter import CharacterTextSplitter
    from langchain.docstore.document import Document
except Exception as e:
    print("必要なライブラリが見つかりません: langchain 等。\n`pip install -r requirements.txt` を実行してください。\nエラー: ", e)
    sys.exit(1)

# LCEL 系のインポート（利用可能なら利用する）
USE_LCEL = True
# 動的 import を使って静的解析エラーの発生を抑え、実行時に利用可能な実装を探す。
import importlib

Runnable = None
RunnablePassthrough = None
try:
    # まず新しい独立パッケージ名を試す
    mod = importlib.import_module("langchain_experimental")
    Runnable = getattr(mod, "Runnable", None)
    RunnablePassthrough = getattr(mod, "RunnablePassthrough", None)
except Exception:
    try:
        # 次に langchain 内の experimental モジュールを試す
        mod2 = importlib.import_module("langchain.experimental.runnable")
        Runnable = getattr(mod2, "Runnable", None)
        RunnablePassthrough = getattr(mod2, "RunnablePassthrough", None)
    except Exception:
        # 利用不可
        Runnable = None
        RunnablePassthrough = None

if Runnable is None or RunnablePassthrough is None:
    USE_LCEL = False

import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 環境変数から API キーを取得
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    print("環境変数 OPENAI_API_KEY が設定されていません。設定してから再実行してください。")
    sys.exit(1)

# --- 設定値 ---
BASE_DIR = Path(__file__).parent
RULE_DIR = BASE_DIR / "rule"
SPEC_DIR = BASE_DIR / "specification"
CHROMA_DIR = BASE_DIR / "chroma_db"
CHROMA_COLLECTION = "specs"

# Embedding モデル名（ローカルで実行できる軽量モデルを使用）
# Chroma の "ローカルの埋め込み" という要件は、ここではローカル HuggingFace 埋め込みを利用して満たす
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# RAG 検索時に取り出すドキュメント数
TOP_K = 3


def load_rules_from_dir(rule_dir: Path) -> List[Dict[str, Any]]:
    """rule ディレクトリ配下の全 JSON を再帰的に読み込み、ルールをフラットなリストで返す。

    ルール JSON は配列 または {"rules": [...] } の形式に対応。
    再帰的な構造（子ルールを 'children' などで持つ）もフラット化して返す。
    """
    rules: List[Dict[str, Any]] = []
    for path in rule_dir.rglob("*.json"):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            logger.warning(f"ルールファイルを読み込めませんでした: {path} - {e}")
            continue

        # 入れ子対応: data が配列か辞書か
        candidates = []
        if isinstance(data, list):
            candidates = data
        elif isinstance(data, dict):
            if "rules" in data and isinstance(data["rules"], list):
                candidates = data["rules"]
            else:
                # 辞書そのものを一つのルール集合とみなす
                candidates = [data]

        # 再帰的にフラット化
        def walk(item: Dict[str, Any], parent_path: str = ""):
            # ルールIDの生成（存在すればそれを利用）
            rid = item.get("id") or item.get("rule_id") or item.get("name") or None
            title = item.get("title") or item.get("name") or rid or "unnamed"
            path_label = f"{parent_path}/{title}" if parent_path else title
            entry = {
                "id": rid,
                "title": title,
                "path": path_label,
                "content": item.get("content") or item.get("description") or json.dumps(item, ensure_ascii=False),
                "raw": item,
            }
            rules.append(entry)
            # 子供があれば再帰
            for key in ("children", "rules", "subrules", "items"):
                if key in item and isinstance(item[key], list):
                    for child in item[key]:
                        if isinstance(child, dict):
                            walk(child, path_label)

        for it in candidates:
            if isinstance(it, dict):
                walk(it)

    logger.info(f"読み込んだルール数: {len(rules)}")
    return rules


def text_from_pdf(path: Path) -> str:
    """シンプルな PDF テキスト抽出。pypdf を利用。ページごとに連結する。"""
    try:
        import pypdf
    except Exception:
        raise RuntimeError("pypdf が必要です。pip install pypdf を実行してください。")
    text_parts = []
    try:
        reader = pypdf.PdfReader(str(path))
        for p in reader.pages:
            txt = p.extract_text() or ""
            text_parts.append(txt)
    except Exception as e:
        logger.warning(f"PDF 読み込み失敗 {path}: {e}")
    return "\n".join(text_parts)


def text_from_docx(path: Path) -> str:
    try:
        import docx
    except Exception:
        raise RuntimeError("python-docx が必要です。pip install python-docx を実行してください。")
    try:
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.warning(f"DOCX 読み込み失敗 {path}: {e}")
        return ""


def load_spec_documents(spec_dir: Path) -> List[Document]:
    """`specification/` 配下のドキュメントを読み込み、langchain Document のリストを返す。

    対応: pdf, docx, md, txt
    メモリ節約: ファイル毎にチャンク分割を行い、最低限のメタデータを付与
    """
    docs: List[Document] = []
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    for path in spec_dir.rglob("*"):
        if path.is_dir():
            continue
        lower = path.suffix.lower()
        try:
            if lower == ".pdf":
                text = text_from_pdf(path)
            elif lower == ".docx":
                text = text_from_docx(path)
            elif lower in (".md", ".txt"):
                text = path.read_text(encoding="utf-8", errors="ignore")
            else:
                # 未対応ファイルは無視
                continue
        except Exception as e:
            logger.warning(f"ファイル読み込み失敗 {path}: {e}")
            continue

        if not text.strip():
            continue

        chunks = text_splitter.split_text(text)
        for i, c in enumerate(chunks):
            meta = {"source": str(path), "chunk": i}
            docs.append(Document(page_content=c, metadata=meta))

    logger.info(f"読み込んだドキュメントチャンク数: {len(docs)}")
    return docs


def init_chroma(docs: List[Document]) -> Chroma:
    """ChromaDB を初期化して、ドキュメントを格納する。既存コレクションがあれば再利用。

    埋め込みはローカル HuggingFace モデルを使う（軽量モデル推奨）
    """
    # HuggingFace 埋め込みラッパー
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

    # Chroma を初期化（persist_directory を指定して永続化）
    vectordb = Chroma(persist_directory=str(CHROMA_DIR), collection_name=CHROMA_COLLECTION, embedding_function=embeddings)

    # 既存が空の場合は追加
    try:
        # Chroma が提供する API により既存数を確認
        existing = vectordb._collection.count() if hasattr(vectordb, "_collection") else None
    except Exception:
        existing = None

    if existing in (None, 0):
        if docs:
            logger.info("Chroma にドキュメントを追加します...")
            vectordb.add_documents(docs)
            vectordb.persist()
    else:
        logger.info("既存の Chroma コレクションを利用します。新規追加は行いません。")

    return vectordb


# --- LCEL ベースの簡易 Wrapper（利用可能なら Runnable を用いる） ---
def make_chat_model() -> ChatOpenAI:
    """ChatOpenAI を作成。モデル名は gpt-4o を指定する。"""
    # ChatOpenAI は環境変数 OPENAI_API_KEY を参照する
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0.0)
    return llm


def agent_a_summarize(llm: ChatOpenAI, rule_text: str, docs: List[Document]) -> str:
    """Agent A: ドキュメント要約エージェント

    - 複数ドキュメントとルールを受け取り、要約を返す
    - 完全性を重視する（重要事項の漏れがないように指示）
    """
    # 取得テキストを簡潔にまとめる（上位 k 件のみを渡す）
    context = "\n\n---関連ドキュメント---\n"
    for d in docs[:TOP_K]:
        src = d.metadata.get("source") if d.metadata else "<unknown>"
        context += f"[source: {src}]\n{d.page_content}\n\n"

    system_prompt = (
        "あなたは優秀なドキュメント要約者です。以下のセキュリティルールを読み、"
        "関連ドキュメントの内容を完全性を保って要約してください。重要な条件、要件、及び検証ポイントを箇条書きで示してください。"
        "\n\n重要: 回答は必ず日本語で行ってください。出力の本文は日本語で記載し、読みやすい箇条書きを心がけてください。"
    )
    messages = [SystemMessage(content=system_prompt), HumanMessage(content=f"ルール:\n{rule_text}\n\nドキュメントコンテキスト:\n{context}")]
    resp = llm(messages)
    return resp.content


def agent_b_check(llm: ChatOpenAI, rule_summary: str, rule_raw: Dict[str, Any], docs: List[Document]) -> Dict[str, Any]:
    """Agent B: ドキュメント確認エージェント

    - ルールに対して、システム（docs）が従っているか評価する
    - 出力は構造化 JSON で返す（result, evidence, details）
    """
    context = "\n\n".join([f"[src:{d.metadata.get('source')}]\n{d.page_content}" for d in docs[:TOP_K]])

    # 厳密な JSON 出力を促すプロンプト（スキーマと例を明示）
    strict_prompt = """
あなたは技術的な評価者です。以下のルール要約と元ルールを読み、与えられたシステム情報がそのルールに従っているかを評価してください。

出力は厳密な JSON のみを返してください。余計な説明や追加テキストは一切書かず、必ず純粋な JSON テキストだけを返してください（コードフェンスや説明を含めないでください）。

JSON スキーマ例:
{
    "result": "〇|△|×",            // 判定
    "evidence": [                   // 推奨: 配列形式
        {"source": "ファイル名や識別子", "excerpt": "抜粋テキスト..."}
    ],
    "details": "追加の説明(任意)"
}

重要: JSON のキー名は英語のままにし、値や説明文は日本語で記載してください。
"""

    init_human = HumanMessage(content=f"ルール要約:\n{rule_summary}\n\n元ルール(raw):\n{json.dumps(rule_raw, ensure_ascii=False)}\n\nドキュメントコンテキスト:\n{context}")

    messages = [SystemMessage(content=strict_prompt), init_human]

    # 最初の回答を取得
    resp = llm(messages)
    # モデルに JSON 形式で答えるよう指示したが、念のためパースを試みる
    text = resp.content

    # デバッグ用: 生出力をログに残す（短縮版）
    logger.debug("Agent B raw output (head 1000 chars): %s", text[:1000].replace('\n', '\\n'))

    # 万が一のため、失敗した生出力をファイルへ追記するユーティリティ
    def _save_model_output(rule_id: str, content: str):
        try:
            logs_dir = BASE_DIR / "logs"
            logs_dir.mkdir(exist_ok=True)
            ts = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
            fname = logs_dir / f"agent_b_output_{ts}_{str(rule_id)[:60].replace(' ', '_')}.log"
            with open(fname, "w", encoding="utf-8") as lf:
                lf.write("--- RAW MODEL OUTPUT ---\n")
                lf.write(content)
            logger.info("モデル出力をログに保存しました: %s", fname)
        except Exception as e:
            logger.debug("モデル出力ログ保存に失敗しました: %s", e)

    # 保存は任意（環境変数で無効化可能）
    if os.environ.get("SAVE_MODEL_OUTPUT", "1") != "0":
        try:
            _save_model_output(rule_raw.get("id") or rule_raw.get("title") or "unknown", text)
        except Exception:
            pass

    parsed = None
    # まず素直に JSON としてデコードを試みる
    try:
        parsed = json.loads(text)
    except Exception:
        parsed = None

    # もしパース失敗したらモデルに再試行を促す（最大2回）
    retries = 0
    while parsed is None and retries < 2:
        retries += 1
        logger.info("JSON パース失敗: モデルへ再試行を行います (試行 %d)。", retries)
        followup = (
            "前の回答は有用でしたが、要求された通り厳密な JSON のみで出力されていませんでした。"
            "以下の JSON スキーマに厳密に合わせ、純粋な JSON テキストのみを出力してください。"
            "\n\nスキーマ: {\"result\":\"〇|△|×\", \"evidence\": [ {\"source\":..., \"excerpt\":...} ], \"details\": \"任意の文字列\" }"
            "\n\n元の出力を参照して、上記スキーマにマッピングして JSON のみを返してください。"
        )
        follow_messages = [SystemMessage(content=strict_prompt), init_human, HumanMessage(content=followup + "\n\n前の出力:\n" + text)]
        try:
            resp2 = llm(follow_messages)
            text2 = resp2.content
            logger.debug("Agent B retry raw output (head 1000): %s", text2[:1000].replace('\n', '\\n'))
            # まず素直に JSON としてデコードを試みる
            try:
                parsed = json.loads(text2)
                text = text2
                break
            except Exception:
                # 次に波括弧ブロックを抽出
                m2 = re.search(r"(\{[\s\S]*\})", text2)
                if m2:
                    try:
                        parsed = json.loads(m2.group(1))
                        text = m2.group(1)
                        break
                    except Exception:
                        parsed = None
                # 最後に簡易変換を試す
                t2 = text2.replace("'", '"')
                t2 = re.sub(r",\s*([}\]])", r"\1", t2)
                t2 = re.sub(r'([\{,\s])(\w+)\s*:', r'\1"\2":', t2)
                try:
                    parsed = json.loads(t2)
                    text = t2
                    break
                except Exception:
                    parsed = None
        except Exception as e:
            logger.debug("モデル再試行中に例外: %s", e)
            parsed = None

    # ここまでで parsed が None ならオリジナル text を用いてフォールバック処理へ
    if parsed is None:
        # JSON の部分文字列を抜き出す試み
        # 1) 最初の波括弧ブロックを抽出して試す
        m = re.search(r"(\{[\s\S]*\})", text)
        if m:
            candidate = m.group(1)
            try:
                parsed = json.loads(candidate)
            except Exception:
                parsed = None

        # 2) シングルクォートをダブルクォートに変換、末尾の余分なカンマを削除、未引用キーに引用付与を試す
        if parsed is None:
            t2 = text.replace("'", '"')
            t2 = re.sub(r",\s*([}\]])", r"\1", t2)
            t2 = re.sub(r'([\{,\s])(\w+)\s*:', r'\1"\2":', t2)
            try:
                parsed = json.loads(t2)
            except Exception:
                parsed = None

    # 最終的にパースできなければヒューリスティック抽出へ
    if parsed is None:
        logger.warning("モデルの出力を JSON としてパースできませんでした。ヒューリスティック抽出を試みます。")

        def _heuristic_parse(text: str) -> Dict[str, Any]:
            out: Dict[str, Any] = {}
            # result (期待値: 〇/△/× または O/X など)
            m = re.search(r"['\"]?result['\"]?\s*[:：]\s*['\"]?([^\"',}\n\r]+)", text, re.IGNORECASE)
            if m:
                out["result"] = m.group(1).strip().strip('"\'')
            else:
                m2 = re.search(r"\b(〇|△|×|O|X|o|x)\b", text)
                if m2:
                    out["result"] = m2.group(1)

            # evidence: try to extract the value after evidence key (allow multiline)
            m_e = re.search(r"['\"]?evidence['\"]?\s*[:：]\s*([\"'])(.*?)\1", text, re.IGNORECASE | re.DOTALL)
            if m_e:
                out["evidence"] = m_e.group(2).strip()
            else:
                m_e2 = re.search(r"evidence\s*[:：\-]\s*(.+)$", text, re.IGNORECASE | re.DOTALL)
                if m_e2:
                    out["evidence"] = m_e2.group(1).strip()

            # details
            m_d = re.search(r"['\"]?details['\"]?\s*[:：]\s*([\"'])(.*?)\1", text, re.IGNORECASE | re.DOTALL)
            if m_d:
                out["details"] = m_d.group(2).strip()
            else:
                m_d2 = re.search(r"details\s*[:：\-]\s*(.+)$", text, re.IGNORECASE | re.DOTALL)
                if m_d2:
                    out["details"] = m_d2.group(1).strip()

            # もし何も抽出できなければ全体を evidence として格納
            if not out.get("evidence") and text:
                out["evidence"] = text.strip()

            # 最低限の result を設定
            if "result" not in out:
                out["result"] = "△"

            return out

        parsed = _heuristic_parse(text)

    # 正規化: evidence は list[ {source, excerpt} ] の形にする
    def _build_evidence_list(evidence_field, docs_list: List[Document]):
        evs = []
        # もしモデルが文字列を返してきたら、そのまま root evidence として docs の抜粋を付与
        if not evidence_field:
            # フォールバック: docs の先頭から抜粋を作る
            for d in docs_list[:TOP_K]:
                evs.append({"source": d.metadata.get("source"), "excerpt": d.page_content[:400].strip()})
            return evs

        if isinstance(evidence_field, str):
            # モデルのフリーテキストをそのまま一つの根拠とする
            evs.append({"source": "(model-output)", "excerpt": evidence_field})
            # さらに docs から抜粋を付与
            for d in docs_list[:TOP_K]:
                evs.append({"source": d.metadata.get("source"), "excerpt": d.page_content[:300].strip()})
            return evs

        # リスト形式が期待される場合
        if isinstance(evidence_field, list):
            for item in evidence_field:
                if isinstance(item, dict):
                    src = item.get("source") or item.get("file") or item.get("path") or "(unknown)"
                    exc = item.get("excerpt") or item.get("text") or json.dumps(item, ensure_ascii=False)
                    evs.append({"source": src, "excerpt": exc[:400].strip()})
                else:
                    evs.append({"source": "(model-output)", "excerpt": str(item)[:400]})
            return evs

        # それ以外の型は文字列化して格納
        evs.append({"source": "(model-output)", "excerpt": str(evidence_field)[:400]})
        return evs

    parsed_evidence = _build_evidence_list(parsed.get("evidence"), docs)
    parsed["evidence_normalized"] = parsed_evidence
    return parsed


# Agent C はユーザの要望により本実装では除外しました。


# Agent D (ダブルチェック) は要件から除外されたため定義は削除しました。


def retrieve_related_docs(vectordb: Chroma, query: str, k: int = TOP_K) -> List[Document]:
    """簡易 RAG 用: query を埋め込み検索し、上位 k のドキュメントを返す"""
    retriever = vectordb.as_retriever(search_kwargs={"k": k})
    results = retriever.get_relevant_documents(query)
    return results


def format_b_result(b_result: Dict[str, Any]) -> str:
    """Agent B の構造化結果を日本語の整形テキストに変換する"""
    lines: List[str] = []
    res = b_result.get("result") or b_result.get("status") or "△"
    lines.append(f"判定: {res}")
    # 詳細説明
    details = b_result.get("details") or b_result.get("detail") or b_result.get("notes")
    if details:
        lines.append("\n説明:")
        if isinstance(details, str):
            lines.append(details)
        else:
            lines.append(json.dumps(details, ensure_ascii=False, indent=2))

    # 根拠（正規化済み）
    evs = b_result.get("evidence_normalized") or []
    if evs:
        lines.append("\n根拠 (参照文書と抜粋):")
        for i, e in enumerate(evs, 1):
            src = e.get("source") or "(unknown)"
            excerpt = e.get("excerpt") or ""
            # 抜粋は改行削除して短めに
            excerpt_clean = excerpt.replace("\n", " ").strip()
            if len(excerpt_clean) > 1200:
                excerpt_clean = excerpt_clean[:1200].rstrip() + " ..."
            lines.append(f"  {i}. source: {src}")
            # indent excerpt
            wrapped = textwrap.fill(excerpt_clean, width=100, subsequent_indent='     ')
            lines.append(textwrap.indent(wrapped, '     '))

    # モデル本体の自由テキスト evidence がある場合
    if b_result.get("evidence") and not evs:
        lines.append("\nモデル出力（根拠）:")
        lines.append(str(b_result.get("evidence")))

    return "\n".join(lines)


 


def find_rule_by_query(rules: List[Dict[str, Any]], query: str) -> Optional[Dict[str, Any]]:
    """ルール一覧から query を元にルールを検索する。ID/パス/タイトル/本文の部分一致で最初のマッチを返す。"""
    q = query.strip().lower()
    # まず ID にマッチ
    for r in rules:
        if r.get("id") and str(r.get("id")).lower() == q:
            return r
    # その他のフィールドで部分マッチ
    for r in rules:
        if q in (r.get("title") or "").lower() or q in (r.get("path") or "").lower() or q in (r.get("content") or "").lower():
            return r
    return None


def interactive_loop(rules: List[Dict[str, Any]], vectordb: Chroma):
    """CUI ベースの簡易チャットループ"""
    llm = make_chat_model()

    help_text = (
        "コマンド一覧:\n"
        "  help                      ヘルプ表示\n"
        "  list                      読み込んだルール一覧の一部を表示\n"
        "  show <query>              ルールを表示（id/title の部分一致）\n"
        "  check <query>             指定したルールに対してシステムが従っているか評価（A->B の順）\n"
        "  showfull <summary|b>      直近のチェックで保存された項目の全文表示\n"
        "  ask <自由テキスト>        システム情報に関する RAG 質問\n"
        "  quit                      終了\n"
    )

    print("AgentRAG チャットボット (CUI)。help と入力してください。\n")
    # 直近の run で作られた出力を保存（全文表示用）
    last_store: Dict[str, Any] = {"summary": None, "b": None}

    def print_section(title: str, content: str, max_len: int = 1200):
        """見やすいセクション表示。

        - セクション見出しを強調表示
        - テキストは段落単位で折り返しを行い、読みやすくする
        - JSON 文字列らしい場合は折り返しせずそのまま出力する
        """
        sep = "=" * 80
        print("\n" + sep)
        print(f"{title}")
        print(sep)
        if content is None:
            print("(なし)\n")
            return

        # プレビュー長を超える場合は省略の目印を付与
        display_text = content
        if isinstance(content, str) and len(content) > max_len:
            display_text = content[:max_len].rstrip() + "\n...（全文は 'showfull' コマンドで表示可）"

        # JSONらしい出力はそのまま表示
        if isinstance(display_text, str) and display_text.strip().startswith(("{", "[")):
            print(display_text)
        else:
            # 段落ごとに折り返して表示（空行で段落分割）
            if isinstance(display_text, str):
                paras = [p.strip() for p in display_text.split("\n\n") if p.strip()]
                for p in paras:
                    wrapped = textwrap.fill(p, width=100)
                    print(wrapped)
                    print()
            else:
                print(str(display_text))
        print(sep + "\n")
    while True:
        try:
            cmd = input("> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n終了します。")
            break

        if not cmd:
            continue
        if cmd == "help":
            print(help_text)
            continue
        if cmd == "list":
            for i, r in enumerate(rules[:50], 1):
                print(f"{i}. id={r.get('id')} title={r.get('title')} path={r.get('path')}")
            continue
        if cmd.startswith("show "):
            q = cmd[len("show "):].strip()
            r = find_rule_by_query(rules, q)
            if not r:
                print("ルールが見つかりませんでした。部分文字列で検索してみてください。")
            else:
                print("--- ルール ---")
                print(f"id: {r.get('id')}")
                print(f"title: {r.get('title')}")
                print(f"path: {r.get('path')}")
                print("content:")
                print(r.get("content"))
            continue

        if cmd.startswith("check "):
            q = cmd[len("check "):].strip()
            r = find_rule_by_query(rules, q)
            if not r:
                print("ルールが見つかりません。別のクエリを試してください。\n(例: ルールの一部の語句や id を入力)\n")
                continue

            print(f"選択されたルール: {r.get('title')} (path: {r.get('path')})")

            # Agent A: 要約（関連ドキュメントを検索して渡す）
            rule_text = r.get("content") or ""
            related = retrieve_related_docs(vectordb, rule_text, k=TOP_K)
            print("[Agent A] ルールと関連ドキュメントから要約を作成しています...")
            summary = agent_a_summarize(llm, rule_text, related)
            last_store["summary"] = summary
            print_section("Agent A - 要約プレビュー", summary)

            # Agent B: 確認
            print("[Agent B] ドキュメントがルールに従っているか評価しています...")
            b_result = agent_b_check(llm, summary, r.get("raw", {}), related)
            last_store["b"] = b_result
            # B の判定表示: プレビュー + 根拠の参照文書表記
            # 整形テキストで表示する
            b_preview_text = format_b_result(b_result)
            print_section("Agent B - 判定（プレビュー）", b_preview_text)
            # 根拠一覧を見やすく表示
            evs = b_result.get("evidence_normalized") or []
            if evs:
                print("根拠 (参照文書と抜粋):")
                for i, e in enumerate(evs, 1):
                    src = e.get("source") or "(unknown)"
                    excerpt = e.get("excerpt") or ""
                    print(f"  {i}. source: {src}")
                    # 抜粋は改行を整形して折り返し表示
                    excerpt_clean = excerpt.replace("\n", " ").strip()
                    if len(excerpt_clean) > 1000:
                        excerpt_clean = excerpt_clean[:1000].rstrip() + " ..."
                    wrapped = textwrap.fill(excerpt_clean, width=100, subsequent_indent='     ')
                    print(textwrap.indent(wrapped, '     '))
                    print()
            else:
                print("(根拠情報はありません)")
            # 補足: 簡易のアクション提案
            result_symbol = b_result.get("result", "△")
            if result_symbol == "〇" or result_symbol == "O" or result_symbol == "o":
                print("補足: 判定は '従っている' と見なされます。必要に応じて関連資料を参照してください。\n")
            elif result_symbol == "×" or result_symbol == "X" or result_symbol == "x":
                print("補足: 判定は '従っていない' です。優先的な対応（修正／設定変更等）が必要です。詳細は関連資料を参照してください。\n")
            else:
                print("補足: 判定は '△'（追加確認が必要）です。関連箇所のログや設定ファイルを追加で提供してください。\n")

            # 評価フロー完了（Agent C は除外）
            print("評価フローが完了しました。必要に応じて 'showfull summary' や 'showfull b' で全文を表示できます。")
            continue

        if cmd.startswith("showfull "):
            what = cmd[len("showfull "):].strip()
            if what not in ("summary", "b"):
                print("'showfull' の引数は summary|b のいずれかを指定してください。")
                continue
            val = last_store.get(what)
            if val is None:
                print(f"まだ '{what}' の出力がありません。先に 'check <query>' を実行してください。")
                continue
            title_map = {"summary": "Agent A - 要約（全文）", "b": "Agent B - 判定（全文）"}
            # dict の場合は整形テキストに変換して出力
            if isinstance(val, str):
                content = val
            else:
                if what == "b":
                    content = format_b_result(val)
                else:
                    content = json.dumps(val, ensure_ascii=False, indent=2)
            print_section(title_map.get(what, what), content, max_len=10_000)
            continue

        if cmd.startswith("ask "):
            q = cmd[len("ask "):].strip()
            # RAG 質問: ドキュメントから上位 TOP_K を引いて LLM に渡す
            docs = retrieve_related_docs(vectordb, q, k=TOP_K)
            context = "\n\n".join([f"[src:{d.metadata.get('source')}]\n{d.page_content}" for d in docs])
            system = "あなたはシステム情報の検索アシスタントです。ユーザの質問に、関連するドキュメントを参照して簡潔に答えてください。"
            messages = [SystemMessage(content=system), HumanMessage(content=f"質問: {q}\n\n参照文書:\n{context}")]
            resp = llm(messages)
            print(resp.content)
            continue

        if cmd in ("quit", "exit", "q"):
            print("終了します。")
            break

        print("不明なコマンドです。help を表示してください。")


def main():
    # ルール読み込み
    rules = load_rules_from_dir(RULE_DIR)

    # ドキュメント読み込み
    docs = load_spec_documents(SPEC_DIR)

    # Chroma 初期化・インデックス作成
    vectordb = init_chroma(docs)

    # インタラクティブループ開始
    interactive_loop(rules, vectordb)


if __name__ == "__main__":
    main()
